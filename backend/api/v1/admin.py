"""管理后台 API（知识库管理完整版 + 原路径修复）"""
import os
import shutil
from typing import List

from fastapi import APIRouter, Depends, UploadFile, File, HTTPException, BackgroundTasks
from sqlalchemy.ext.asyncio import AsyncSession

from models.db import get_db
from models.user import User
from api.deps import get_current_user
from utils.logger_handler import logger
from schemas.common import ResponseBase
from utils.config_handler import chroma_conf      # 【关键】和 vector_store 共用配置
from utils.path_tool import get_abs_path          # 【关键】和 vector_store 共用路径解析

router = APIRouter()

# 【关键修复】UPLOAD_DIR 改为和 vector_store.py 完全一致
# 原来硬编码的 "data/uploads" 改为 chroma_conf["data_path"]，即 backend/data/
UPLOAD_DIR = get_abs_path(chroma_conf.get("data_path", "data"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

logger.info(f"[Admin] 知识库目录: {UPLOAD_DIR}")


@router.get("/dashboard/stats")
async def get_dashboard_stats(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """仪表盘统计数据"""
    from models.chat import ChatSession, ChatMessage
    from sqlalchemy import func, select

    total_sessions = await db.scalar(select(func.count(ChatSession.id)))
    total_messages = await db.scalar(select(func.count(ChatMessage.id)))

    # 扫描实际文件（和 vector_store 同一目录）
    file_count = 0
    if os.path.exists(UPLOAD_DIR):
        file_count = len([
            f for f in os.listdir(UPLOAD_DIR)
            if os.path.isfile(os.path.join(UPLOAD_DIR, f))
            and f.lower().endswith(('.pdf', '.txt', '.docx', '.md'))
        ])

    logger.info(f"[Admin] 扫描知识库: {UPLOAD_DIR}, 找到 {file_count} 个文件")

    return {
        "success": True,
        "data": {
            "total_users": 0,
            "total_sessions": total_sessions or 0,
            "total_messages": total_messages or 0,
            "knowledge_files": file_count
        }
    }


@router.post("/knowledge/upload")
async def upload_knowledge(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """上传知识库文件到原 data/ 目录"""
    allowed_extensions = {".pdf", ".txt", ".docx", ".md"}
    ext = os.path.splitext(file.filename)[1].lower()

    if ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，仅支持: {', '.join(allowed_extensions)}"
        )

    safe_filename = os.path.basename(file.filename)
    if ".." in safe_filename or "/" in safe_filename or "\\" in safe_filename:
        raise HTTPException(status_code=400, detail="非法文件名")

    # 保存到和 vector_store 同一目录
    file_path = os.path.join(UPLOAD_DIR, safe_filename)

    # 如果文件已存在，先删旧向量
    if os.path.exists(file_path):
        try:
            _delete_vector_by_source(file_path)
        except Exception as e:
            logger.warning(f"删除旧向量失败（可能不存在）: {e}")

    try:
        # 1. 保存文件
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 2. 读取内容
        content = _read_file_content(file_path, ext)
        if not content or len(content.strip()) == 0:
            os.remove(file_path)
            raise HTTPException(status_code=400, detail="文件内容为空")

        # 3. 分割文本
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from langchain_core.documents import Document

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", "！", "？", " ", ""]
        )
        chunks = splitter.split_text(content)

        # 4. 构建 Document 列表
        documents = [
            Document(
                page_content=chunk,
                metadata={"source": file_path, "filename": safe_filename, "chunk_index": i}
            )
            for i, chunk in enumerate(chunks)
        ]

        # 5. 存入向量库（使用带用量追踪的接口）
        from rag.vector_store import VectorStoreService
        vs = VectorStoreService()
        vs.add_documents_with_tracking(documents)

        logger.info(f"知识库上传成功: {safe_filename}, 共 {len(chunks)} 个片段")

        return {
            "success": True,
            "message": "上传成功",
            "data": {
                "filename": safe_filename,
                "chunks": len(chunks),
                "size": os.path.getsize(file_path)
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        logger.error(f"知识库上传失败: {e}")
        raise HTTPException(status_code=500, detail=f"上传失败: {str(e)}")


@router.post("/knowledge/batch-upload")
async def batch_upload_knowledge(
    files: List[UploadFile] = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """批量上传知识库文件"""
    results = []
    for file in files:
        try:
            result = await upload_knowledge(file, current_user, db)
            results.append({"filename": file.filename, "status": "success", "data": result.get("data")})
        except HTTPException as e:
            results.append({"filename": file.filename, "status": "error", "message": e.detail})
        except Exception as e:
            results.append({"filename": file.filename, "status": "error", "message": str(e)})

    success_count = sum(1 for r in results if r["status"] == "success")
    return {
        "success": True,
        "message": f"批量上传完成: {success_count}/{len(files)} 成功",
        "data": results
    }


@router.post("/knowledge/rebuild")
async def rebuild_knowledge_base(
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """重建整个向量库"""
    background_tasks.add_task(_rebuild_vector_store_task)
    return {
        "success": True,
        "message": "向量库重建任务已启动，请稍后刷新查看"
    }


@router.get("/knowledge/list")
async def list_knowledge(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """获取已上传的知识库文件列表（扫描原 data/ 目录）"""
    files = []

    logger.info(f"[Admin] 列出知识库: {UPLOAD_DIR}, exists={os.path.exists(UPLOAD_DIR)}")

    if not os.path.exists(UPLOAD_DIR):
        return {"success": True, "data": []}

    for filename in os.listdir(UPLOAD_DIR):
        file_path = os.path.join(UPLOAD_DIR, filename)
        if os.path.isfile(file_path) and filename.lower().endswith(('.pdf', '.txt', '.docx', '.md')):
            stat = os.stat(file_path)
            files.append({
                "id": filename,
                "filename": filename,
                "size": stat.st_size,
                "uploaded_at": stat.st_mtime,
                "path": file_path
            })

    files.sort(key=lambda x: x["uploaded_at"], reverse=True)
    logger.info(f"[Admin] 返回 {len(files)} 个知识库文件")
    return {"success": True, "data": files}


@router.delete("/knowledge/{filename}")
async def delete_knowledge(
    filename: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """删除知识库文件（同时清向量库）"""
    safe_filename = os.path.basename(filename)
    if ".." in safe_filename or "/" in safe_filename or "\\" in safe_filename:
        raise HTTPException(status_code=400, detail="非法文件名")

    file_path = os.path.join(UPLOAD_DIR, safe_filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")

    try:
        _delete_vector_by_source(file_path)
        os.remove(file_path)
        return {"success": True, "message": "删除成功"}
    except Exception as e:
        logger.error(f"删除失败: {e}")
        raise HTTPException(status_code=500, detail=f"删除失败: {str(e)}")


# ========== 辅助函数 ==========

def _read_file_content(file_path: str, ext: str) -> str:
    """读取不同格式文件"""
    if ext in (".txt", ".md"):
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError("无法解码文本文件")

    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            return "".join([page.extract_text() or "" for page in reader.pages])
        except ImportError:
            raise HTTPException(status_code=500, detail="缺少 pypdf，请执行: pip install pypdf")

    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        except ImportError:
            raise HTTPException(status_code=500, detail="缺少 python-docx，请执行: pip install python-docx")

    else:
        raise ValueError(f"不支持的格式: {ext}")


def _delete_vector_by_source(file_path: str):
    """按 source 元数据删除向量"""
    try:
        from rag.vector_store import VectorStoreService
        vs = VectorStoreService()
        if hasattr(vs, 'vector_store') and hasattr(vs.vector_store, '_collection'):
            vs.vector_store._collection.delete(where={"source": file_path})
            logger.info(f"[VectorStore] 已删除来源为 {file_path} 的向量")
    except Exception as e:
        logger.warning(f"向量删除失败（可能不存在）: {e}")


def _rebuild_vector_store_task():
    """后台任务：重建向量库"""
    try:
        logger.info("[VectorStore] 开始重建向量库...")

        from rag.vector_store import VectorStoreService
        vs = VectorStoreService()
        if hasattr(vs, 'vector_store') and hasattr(vs.vector_store, '_collection'):
            vs.vector_store._collection.delete(where={})
            logger.info("[VectorStore] 已清空现有向量库")

        vs.load_document()

        logger.info("[VectorStore] 向量库重建完成")
    except Exception as e:
        logger.error(f"[VectorStore] 重建失败: {e}", exc_info=True)